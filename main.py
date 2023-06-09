import difflib
import fire
import json
import os
import shutil
import subprocess
import sys
import re
import argparse
import openai
import tiktoken
from termcolor import cprint
from dotenv import load_dotenv
from docx import Document

load_dotenv()

parser = argparse.ArgumentParser()

parser.add_argument("file", help="Path of the file to be translated")
parser.add_argument("-m", "--model", default="gpt-3.5-turbo", help="GPT model")
parser.add_argument("-lang", "--target-language", default="english", help="Target language to translate to")

args = parser.parse_args()

openai.api_key = os.getenv("OPENAI_API_KEY")
MAX_TOKENS_MODEL = int(os.getenv("MAX_TOKENS_MODEL"))
ENCODER = tiktoken.encoding_for_model(args.model)

with open("auto-translator-prompt.txt") as f:
    SYSTEM_PROMPT = f.read()


class AutoTranslator:
    def __init__(
        self, file_path = None, model = None, system_prompt = None, max_tokens_per_request = None, target_language=None,
        auto_correct = False, auto_improve = False
    ):
        self.file_path = file_path
        self.file_extension = file_path.rsplit('.')[1]
        self.translated_file_path = f"{file_path.rsplit('.')[0]}-translated-to-{target_language}.{file_path.rsplit('.')[1]}"
        self.model = model
        self.system_prompt = system_prompt
        self.target_language = target_language
        self.auto_correct = auto_correct
        self.auto_improve = auto_improve
        self.parameters = f"""Parameters:
            - TARGET-LANGUAGE: {self.target_language}
            - AUTO-CORRECT: {self.auto_correct}
            - AUTO-IMPROVE: {self.auto_improve}
        """

        self.lines = []
        self.first_line = 0
        self.last_line = 0
        self.number_of_lines = 0
        self.batch = 0
        self.lines_per_batch = 0
        self.max_tokens_per_request = max_tokens_per_request

        self.raw_text = ""
        self.translated_text = ""

        self.doc = Document()

    def get_amount_of_lines_in_file(self):
        if self.file_extension == "docx":
            self.number_of_lines = len(self.lines)
            return

        with open(self.file_path, "r") as file:
            self.number_of_lines = sum(1 for _ in file)

    def calculate_lines_per_batch(self):
        estimated_char_per_line = 1000
        # 1 token = 4 char
        self.lines_per_batch = int(
            self.max_tokens_per_request / (estimated_char_per_line / 4)
        )
        total_tokens = 4 * estimated_char_per_line * self.number_of_lines
        print(
            f"Estimated amount of tokens is {total_tokens} and will take {self.number_of_lines / self.lines_per_batch}"
        )

    def write_translated_text(self):
        if self.file_extension == "docx":
            translated_lines = dict(re.findall(r"<(\d+)>(.*?)<\/\1>", self.translated_text, flags=re.DOTALL))
            i = self.first_line
            for p in self.doc.paragraphs:
                for run in p.runs:
                    if run.text.strip() and str(i) in translated_lines and re.match(rf"<{i}>(.*?)</{i}>", run.text, flags=re.DOTALL):
                        run.text = translated_lines[str(i)]
                        i += 1

            self.doc.save(self.translated_file_path)
            return

        with open(f"{self.translated_file_path}", "a+") as translated_file:
            translated_file.write(self.translated_text)

    def read_whole_file(self):
        if self.file_extension == "docx":
            self.doc = Document(self.file_path)
            i = 0
            for paragraph in self.doc.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        run.text = f"<{i}>{run.text}</{i}>"
                        self.lines.append(run.text)
                        i += 1
            return

        with open(self.file_path, "r") as file:
            self.lines = file.readlines()

    def send_to_gpt(self):
        messages = [
            {
                "role": "system",
                "content": self.system_prompt,
            },
            {
                "role": "system",
                "content": self.parameters
            },
            {
                "role": "user",
                "content": self.raw_text,
            },
        ]
        response = openai.ChatCompletion.create(
            model=self.model,
            messages=messages,
            temperature=0.5,
        )
        
        self.translated_text = response['choices'][0]['message']['content']
        # return response['choices'][0]['message']['content']

    def translate(self):
        self.read_whole_file()
        self.get_amount_of_lines_in_file()
        self.calculate_lines_per_batch()
        while self.last_line < self.number_of_lines:
            print(f"Processing batch {self.batch}...")
            self.last_line = (self.batch + 1) * self.lines_per_batch
            if self.last_line > self.number_of_lines:
                self.last_line = self.number_of_lines
            print(f"Sending lines {self.first_line} - {self.last_line}")

            self.raw_text = "".join(self.lines[self.first_line:self.last_line])
            self.send_to_gpt()
            self.write_translated_text()
            self.batch += 1
            self.first_line = self.last_line + 1 if self.file_extension != "docx" else self.last_line


translator = AutoTranslator(
    file_path=args.file,
    model=args.model,
    target_language=args.target_language,
    system_prompt=SYSTEM_PROMPT,
    max_tokens_per_request=MAX_TOKENS_MODEL
)
translator.translate()

